# =============================================================================
# @project    ELY — Exactly Like You
# @file       backend/tests/test_pdf_to_docx_structural.py
# @brief      Reconstruire la STRUCTURE LOGIQUE d'un PDF, au lieu d'un
#             paragraphe par ligne visuelle.
# @license    MIT
# =============================================================================
"""Pins de la conversion PDF → DOCX structurelle.

**Ce que faisait Ely.** Le cœur de ``_write_docx`` tenait en trois lignes :

    for line in text.split("\\n"):
        if line.strip():
            document.add_paragraph(line.strip())

Une ligne VISUELLE du PDF = un paragraphe Word. C'est exactement le premier
raccourci que décrit le mode d'emploi de Franck : *« on obtient 12 000 faux
paragraphes, des mots coupés par les césures, aucun style »*. Et
``pypdf.extract_text()`` ne donne aucune géométrie — ni coordonnées, ni corps,
ni graisse — donc aucun des pièges ne pouvait être traité.

**Le principe repris.** Un PDF ne contient pas de paragraphes, il contient des
glyphes posés à des coordonnées. La seule voie utile est de **reconstruire la
structure logique depuis la géométrie**, puis de l'écrire en styles Word
natifs.

**Les pièges couverts par ce lot**, dans l'ordre du mode d'emploi :

- **(a) fragments de ligne** — un changement d'italique crée un bloc distinct,
  parfois décalé de 1 à 2 points. Trié naïvement par ``y``, le texte part dans
  le désordre. On regroupe par tolérance verticale, puis on trie par ``x``.
- **(b) césures** — on supprime le trait d'union si la ligne suivante commence
  par une minuscule ou une apostrophe. Sinon on le garde : sans cette règle on
  obtient « éblouis-sement » ou « portefenêtre ».
- **(c) paragraphes à cheval sur deux pages** — l'erreur la plus visible. On
  recolle quand la dernière ligne de la page remplit la justification et ne se
  termine pas par une ponctuation de fin de phrase.
- **(d) titres** — jamais par le texte (« Chapitre » ne marche pas : langue,
  chiffres romains, titres muets). On croise graisse, corps supérieur au corps
  courant, position haute et nombre de lignes, confirmé par le signet PDF.
- **styles nommés** plutôt que formatage direct : modifier le corps du roman
  doit être UNE modification, pas quatre mille.
- **contrôle 0** — comparer le DOCX au modèle ne prouve rien si le modèle est
  déjà amputé. On compare le multi-ensemble des caractères PDF ↔ DOCX.

**Calibrage.** Les constantes ne sont pas universelles : ce sont les mesures de
CE document. Elles sont donc relevées sur le PDF lui-même (pics d'avance
verticale), jamais devinées.

**Hors périmètre de ce lot, et assumé** : vers et blocs centrés, images en
ligne, harmonisation de la ponctuation en fausse fonte, champs PAGEREF du
sommaire. Ce sont les raffinements ; ce lot livre la structure.

Run with:  cd backend && python -m pytest tests/test_pdf_to_docx_structural.py -v
"""
from __future__ import annotations

import pytest

fitz = pytest.importorskip("fitz", reason="pymupdf requis pour ces pins")

# Géométrie du PDF de test : deux pics nets, comme sur un vrai manuscrit.
_LINE_PITCH = 14.0     # interligne courant
_PARA_GAP = 26.0       # passage au paragraphe suivant
_LEFT = 72.0
_TOP = 100.0
_BODY_SIZE = 11.0


def _page_with_lines(doc, lines, *, left=_LEFT, top=_TOP):
    """Pose des lignes à une géométrie EXACTE.

    Chaque entrée : ``(texte, avance_avant, taille, gras)``.
    """
    page = doc.new_page(width=595, height=842)
    y = top
    for text, gap, size, bold in lines:
        y += gap
        font = "hebo" if bold else "helv"
        page.insert_text((left, y), text, fontsize=size, fontname=font)
    return page


def _docx_paragraphs(path):
    import docx

    return [p for p in docx.Document(str(path)).paragraphs if p.text.strip()]


def _convert(tmp_path, doc, name="test"):
    from app.services.pdf_to_docx import convert_pdf_to_docx

    pdf_path = tmp_path / f"{name}.pdf"
    doc.save(str(pdf_path))
    out = tmp_path / f"{name}.docx"
    report = convert_pdf_to_docx(pdf_path.read_bytes(), out)
    return out, report


# ------------------------------------------------- le défaut d'origine

def test_a_flowing_paragraph_is_not_split_into_one_paragraph_per_line(tmp_path):
    """LE test du lot. Quatre lignes d'un même paragraphe doivent donner UN
    paragraphe, pas quatre — c'est tout l'écart entre un manuscrit éditable et
    12 000 faux paragraphes."""
    doc = fitz.open()
    _page_with_lines(doc, [
        ("Le vent se levait sur la ville et personne ne semblait", 0, _BODY_SIZE, False),
        ("y prendre garde, pas meme les oiseaux qui tournaient", _LINE_PITCH, _BODY_SIZE, False),
        ("au-dessus des toits en poussant leurs cris rauques", _LINE_PITCH, _BODY_SIZE, False),
        ("dans la lumiere declinante du soir.", _LINE_PITCH, _BODY_SIZE, False),
    ])
    out, _ = _convert(tmp_path, doc)

    assert len(_docx_paragraphs(out)) == 1


def test_a_wider_gap_starts_a_new_paragraph(tmp_path):
    """Le seuil se place ENTRE les deux pics d'avance verticale, pas au
    milieu : les paragraphes serrés sont plus fréquents que les interlignes
    dilatés."""
    doc = fitz.open()
    _page_with_lines(doc, [
        ("Premiere ligne du premier paragraphe qui occupe la largeur", 0, _BODY_SIZE, False),
        ("et sa suite immediate sur la ligne suivante du meme bloc.", _LINE_PITCH, _BODY_SIZE, False),
        ("Deuxieme paragraphe, separe par une avance plus large ici.", _PARA_GAP, _BODY_SIZE, False),
    ])
    out, _ = _convert(tmp_path, doc)

    assert len(_docx_paragraphs(out)) == 2


# ------------------------------------------------------- piège (b) césures

def test_a_typographic_hyphen_is_removed(tmp_path):
    """« éblouis- / sement » doit redonner « éblouissement » : la ligne
    suivante commence par une minuscule."""
    doc = fitz.open()
    _page_with_lines(doc, [
        ("Elle fut saisie par un immense eblouis-", 0, _BODY_SIZE, False),
        ("sement au moment de franchir la porte.", _LINE_PITCH, _BODY_SIZE, False),
    ])
    out, _ = _convert(tmp_path, doc)

    text = _docx_paragraphs(out)[0].text
    assert "eblouissement" in text
    assert "eblouis-" not in text


def test_a_real_hyphen_is_kept(tmp_path):
    """« porte- / Fenetre » garde son trait d'union : la suite commence par
    une majuscule. Sans cette nuance on recolle « portefenetre »."""
    doc = fitz.open()
    _page_with_lines(doc, [
        ("Il ouvrit la porte-", 0, _BODY_SIZE, False),
        ("Fenetre pour respirer un peu.", _LINE_PITCH, _BODY_SIZE, False),
    ])
    out, _ = _convert(tmp_path, doc)

    assert "porte-Fenetre" in _docx_paragraphs(out)[0].text


# --------------------------------------- piège (c) paragraphes inter-pages

def test_a_paragraph_spanning_a_page_break_is_merged(tmp_path):
    """L'erreur la plus visible : un paragraphe coupé par un saut de page
    devenait deux paragraphes, et le manuscrit se retrouvait avec des
    centaines de fausses ruptures."""
    doc = fitz.open()
    _page_with_lines(doc, [
        ("Le recit continuait sans interruption apparente et la phrase", 0, _BODY_SIZE, False),
        ("se poursuivait jusqu au bas de la page sans jamais s achever", _LINE_PITCH, _BODY_SIZE, False),
    ], top=700)
    _page_with_lines(doc, [
        ("sur cette derniere ligne qui la termine enfin proprement.", 0, _BODY_SIZE, False),
    ])
    out, _ = _convert(tmp_path, doc)

    paras = _docx_paragraphs(out)
    assert len(paras) == 1, f"{len(paras)} paragraphes — la rupture de page n'a pas été recollée"


def test_a_finished_sentence_before_a_page_break_is_not_merged(tmp_path):
    """Garde-fou symétrique : une phrase qui se termine par un point reste un
    paragraphe distinct. Tout recoller serait aussi faux que ne rien recoller."""
    doc = fitz.open()
    _page_with_lines(doc, [
        ("Une phrase complete qui se termine proprement en bas de page.", 0, _BODY_SIZE, False),
    ], top=700)
    _page_with_lines(doc, [
        ("Un nouveau paragraphe commence sur la page suivante.", 0, _BODY_SIZE, False),
    ])
    out, _ = _convert(tmp_path, doc)

    assert len(_docx_paragraphs(out)) == 2


# --------------------------------------------------- piège (d) les titres

def test_a_title_is_detected_by_geometry_not_by_its_words(tmp_path):
    """Chercher « Chapitre » ne marche pas — langue, chiffres romains, titres
    muets. On croise graisse, corps et position."""
    doc = fitz.open()
    _page_with_lines(doc, [
        ("Le seuil", 0, 20.0, True),
        ("Le vent se levait sur la ville et personne n y prenait garde", 40.0, _BODY_SIZE, False),
        ("ce matin la, malgre les nuages bas qui couraient au-dessus", _LINE_PITCH, _BODY_SIZE, False),
        ("des toits sombres et luisants de la pluie de la veille au", _LINE_PITCH, _BODY_SIZE, False),
        ("soir, quand tout le monde dormait encore profondement.", _LINE_PITCH, _BODY_SIZE, False),
        ("Un second paragraphe suivait, plus court mais bien present.", _PARA_GAP, _BODY_SIZE, False),
    ])
    out, report = _convert(tmp_path, doc)

    paras = _docx_paragraphs(out)
    assert paras[0].style.name != paras[1].style.name, (
        "le titre reçoit le même style que le corps — le volet Navigation de "
        "Word restera vide"
    )
    assert report.titles == 1


# ------------------------------------------------------ les styles nommés

def test_the_document_uses_named_styles_not_direct_formatting(tmp_path):
    """`python-docx` pousse au formatage direct, et c'est ce qui produit ces
    DOCX impossibles à retoucher : changer le corps demande 4 000
    modifications. Chaque nature de paragraphe reçoit un style nommé."""
    doc = fitz.open()
    _page_with_lines(doc, [
        ("Le seuil", 0, 20.0, True),
        ("Le vent se levait sur la ville et personne n y prenait garde", 40.0, _BODY_SIZE, False),
        ("ce matin la, malgre les nuages bas qui couraient au-dessus", _LINE_PITCH, _BODY_SIZE, False),
        ("des toits sombres et luisants de la pluie de la veille au", _LINE_PITCH, _BODY_SIZE, False),
    ])
    out, _ = _convert(tmp_path, doc)

    import docx

    styles = {s.name for s in docx.Document(str(out)).styles}
    assert "Chapitre" in styles


def test_the_title_style_feeds_the_navigation_pane(tmp_path):
    """`w:outlineLvl` alimente le volet Navigation de Word et les signets à
    l'export PDF. Sans lui, un titre n'est qu'un texte en gras."""
    doc = fitz.open()
    _page_with_lines(doc, [
        ("Le seuil", 0, 20.0, True),
        ("Du corps de texte ordinaire ici, sur plusieurs lignes pour", 40.0, _BODY_SIZE, False),
        ("que le calibrage ait de quoi mesurer un interligne reel.", _LINE_PITCH, _BODY_SIZE, False),
        ("Une troisieme ligne acheve de rendre l echantillon credible.", _LINE_PITCH, _BODY_SIZE, False),
    ])
    out, _ = _convert(tmp_path, doc)

    import docx

    style = docx.Document(str(out)).styles["Chapitre"]
    assert style.element.xpath(".//w:outlineLvl"), "le style de titre n'a pas de niveau de plan"


# ------------------------------------------------------- contrôle 0

def test_no_character_is_lost_between_the_pdf_and_the_docx(tmp_path):
    """Le contrôle qu'on oublie. Vérifier le DOCX contre un modèle ne prouve
    rien si le modèle est déjà amputé : on compare le multi-ensemble des
    caractères directement au PDF."""
    doc = fitz.open()
    _page_with_lines(doc, [
        ("Le vent se levait sur la ville et personne ne semblait", 0, _BODY_SIZE, False),
        ("y prendre garde ce matin la, malgre les nuages bas.", _LINE_PITCH, _BODY_SIZE, False),
        ("Un second paragraphe, pour faire bonne mesure ici.", _PARA_GAP, _BODY_SIZE, False),
    ])
    out, report = _convert(tmp_path, doc)

    assert report.missing_chars == 0, (
        f"caractères perdus à la conversion : {report.missing_chars}"
    )
    assert report.chars_pdf > 100


def test_the_report_says_what_it_measured(tmp_path):
    """Un rapport qui dit « converti » sans chiffre ne permet pas de savoir si
    la conversion a réussi — c'est la façade que le projet traque partout."""
    doc = fitz.open()
    _page_with_lines(doc, [("Une ligne de texte simple et complete.", 0, _BODY_SIZE, False)])
    _, report = _convert(tmp_path, doc)

    assert report.pages == 1
    assert report.paragraphs >= 1
    assert report.chars_pdf > 0


# ------------------------------------------------------------ robustesse

def test_a_pdf_without_a_text_layer_is_refused_explicitly(tmp_path):
    """Un scan doit donner une erreur claire, jamais un Word vide livré comme
    un succès."""
    from app.services.pdf_to_docx import NoTextLayer, convert_pdf_to_docx

    doc = fitz.open()
    doc.new_page(width=595, height=842)
    pdf_path = tmp_path / "scan.pdf"
    doc.save(str(pdf_path))

    with pytest.raises(NoTextLayer):
        convert_pdf_to_docx(pdf_path.read_bytes(), tmp_path / "scan.docx")


def test_calibration_is_measured_on_the_document(tmp_path):
    """Les constantes ne sont pas universelles : ce sont les mesures de CE
    livre. Un document au pas de 20 pt ne doit pas être découpé avec le seuil
    d'un document au pas de 14 pt."""
    from app.services.pdf_to_docx import calibrate

    doc = fitz.open()
    _page_with_lines(doc, [(f"Ligne numero {i} du meme paragraphe continu ici.",
                            0 if i == 0 else 20.0, _BODY_SIZE, False) for i in range(8)])
    pdf_path = tmp_path / "pitch20.pdf"
    doc.save(str(pdf_path))

    cal = calibrate(fitz.open(str(pdf_path)))

    assert 19.0 < cal.line_pitch < 21.0, f"pas mesuré : {cal.line_pitch}"
    assert cal.para_gap_min > cal.line_pitch


# =============================================================================
# Typographie — ce que la première version extrayait puis JETAIT
# =============================================================================
#
# Franck, sur un manuscrit réel de 395 pages : « aucune taille de texte
# récupérée, en fait elle a extrait le texte brut, l'a copié-collé dans un
# .docx et voilà ». Mesure faite, il avait raison sur tout :
#
#     PDF source              Ce qui sortait
#     396 × 612 pt            612 × 792 (US Letter, défaut python-docx)
#     Times / PublicoText     aucune police
#     12 · 10 · 14 · 22 pt    aucun corps
#     64 spans en italique    0
#     2 couleurs              0
#
# Le pire : ces attributs étaient DÉJÀ LUS span par span pour détecter les
# titres, puis abandonnés au moment d'écrire. Pas un problème difficile — une
# donnée extraite et jetée. Les classer dans les « raffinements » était un
# mauvais jugement : pour un manuscrit, c'est le minimum pour que le document
# se ressemble.

def test_the_page_keeps_the_source_format(tmp_path):
    """LE test de cette reprise. Un livre en 396 × 612 ne doit pas ressortir
    en US Letter — c'est la première chose qu'on voit en ouvrant le fichier."""
    doc = fitz.open()
    page = doc.new_page(width=396, height=612)
    y = 100
    for i in range(6):
        page.insert_text((50, y), f"Ligne {i} d un paragraphe continu du livre.",
                         fontsize=_BODY_SIZE, fontname="helv")
        y += _LINE_PITCH
    out, _ = _convert(tmp_path, doc)

    import docx

    section = docx.Document(str(out)).sections[0]
    assert round(section.page_width.pt) == 396
    assert round(section.page_height.pt) == 612


def test_the_body_font_and_size_are_carried_over(tmp_path):
    """Le corps dominant du document devient celui du style `Normal` — donc
    UNE modification si Franck veut le changer, pas quatre mille."""
    doc = fitz.open()
    page = doc.new_page(width=396, height=612)
    y = 100
    for i in range(8):
        page.insert_text((50, y), f"Ligne {i} en corps de texte du manuscrit.",
                         fontsize=13.0, fontname="tiro")
        y += _LINE_PITCH
    out, _ = _convert(tmp_path, doc)

    import docx

    normal = docx.Document(str(out)).styles["Normal"]
    assert normal.font.size is not None, "aucun corps posé — Word appliquera le sien"
    assert abs(normal.font.size.pt - 13.0) < 0.6
    assert normal.font.name, "aucune police posée"


def test_an_italic_span_becomes_an_italic_run(tmp_path):
    """64 spans en italique sur 55 pages d'un roman : c'est du dialogue, des
    titres d'œuvres, des mots étrangers. Les perdre change le texte."""
    doc = fitz.open()
    page = doc.new_page(width=396, height=612)
    y = 100
    for i in range(5):
        page.insert_text((50, y), f"Ligne ordinaire numero {i} du paragraphe.",
                         fontsize=_BODY_SIZE, fontname="helv")
        y += _LINE_PITCH
    page.insert_text((50, y + _PARA_GAP), "Un mot en italique ici precisement.",
                     fontsize=_BODY_SIZE, fontname="tiit")
    out, _ = _convert(tmp_path, doc)

    runs = [r for p in _docx_paragraphs(out) for r in p.runs]
    assert any(r.italic for r in runs), "toutes les italiques du document ont disparu"


def test_a_size_change_inside_the_text_is_preserved(tmp_path):
    """Le PDF utilise 12, 10, 14 et 22 pt. Tout ramener au corps courant
    écrase les exergues, les notes et les titres de partie."""
    doc = fitz.open()
    page = doc.new_page(width=396, height=612)
    y = 100
    for i in range(6):
        page.insert_text((50, y), f"Corps courant, ligne {i} du paragraphe.",
                         fontsize=12.0, fontname="helv")
        y += _LINE_PITCH
    page.insert_text((50, y + _PARA_GAP), "Une note en petit corps.",
                     fontsize=8.0, fontname="helv")
    out, _ = _convert(tmp_path, doc)

    sizes = {r.font.size.pt for p in _docx_paragraphs(out) for r in p.runs
             if r.font.size is not None}
    assert any(s < 10 for s in sizes), f"corps trouvés : {sizes} — la note a été normalisée"


def test_typography_does_not_break_the_integrity_check(tmp_path):
    """Garde-fou : écrire des runs formatés ne doit pas perdre de caractère en
    route. Le contrôle 0 reste la référence."""
    doc = fitz.open()
    page = doc.new_page(width=396, height=612)
    y = 100
    for i in range(6):
        page.insert_text((50, y), f"Ligne {i} avec du texte a preserver entierement.",
                         fontsize=_BODY_SIZE, fontname="helv")
        y += _LINE_PITCH
    page.insert_text((50, y + _PARA_GAP), "Et une phrase en italique pour finir.",
                     fontsize=_BODY_SIZE, fontname="tiit")
    _, report = _convert(tmp_path, doc)

    assert report.missing_chars == 0
