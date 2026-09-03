# =============================================================================
# @project    ELY — Exactly Like You
# @file       backend/tests/test_pdf_docx_materializes_plan.py
# @brief      Ely matérialise le balisage du modèle — mécaniquement, sans rien
#             réinterpréter, et en chiffrant ce qu'elle a retiré.
# @license    MIT
# =============================================================================
"""La moitié « Ely matérialise et contrôle » du lot.

Le fait qui a ouvert ce lot
----------------------------
    add_page_break dans app/services/pdf_to_docx.py : 0 occurrence

Mesure faite pendant le lot, qui précise la formule : le convertisseur produit
bien des sauts, mais **seulement** par ``w:pageBreakBefore`` posé sur les
styles ``Chapitre``/``Partie``. Un saut n'existe donc que là où l'heuristique a
cru voir un titre. Sur les 12 premières pages du manuscrit de Franck : 12 pages
PDF → 6 sauts, la page de copyright classée « chapitre », la dédicace et
l'exergue recollées au texte. D'où « 4-5 pages aplaties en une ».

Ce que ces pins verrouillent
-----------------------------
- un ``SAUT`` demandé par le modèle produit un VRAI saut de page ;
- un ``ARTEFACT`` (le folio resté en texte courant) sort du document, et le
  nombre de caractères retirés est **dit**, pas fondu dans « caractères
  perdus » — sans quoi le contrôle 0 deviendrait un bruit permanent ;
- une ``SUITE`` recolle un paragraphe à cheval sur deux pages ;
- **un plan vide ne change rien.** C'est la règle « échouer OUVERT » de la
  boucle de conformité : un modèle en panne rend la main au comportement
  d'avant, il ne mutile pas le document.

Run with:  cd backend && python -m pytest tests/test_pdf_docx_materializes_plan.py -v
"""
from __future__ import annotations

import pytest

fitz = pytest.importorskip("fitz", reason="pymupdf requis pour ces pins")

_LINE_PITCH = 14.0
_PARA_GAP = 26.0
_LEFT = 72.0
_BODY_SIZE = 11.0


def _page(doc, lines, *, left=_LEFT, top=100.0):
    page = doc.new_page(width=595, height=842)
    y = top
    for text, gap, size, bold in lines:
        y += gap
        page.insert_text((left, y), text, fontsize=size,
                         fontname="hebo" if bold else "helv")
    return page


def _body(n, prefix="Ligne"):
    return [(f"{prefix} {i} du paragraphe courant qui coule sans fin.",
             0 if i == 0 else _LINE_PITCH, _BODY_SIZE, False) for i in range(n)]


def _convert(tmp_path, doc, plan=None, name="plan"):
    from app.services.pdf_to_docx import convert_pdf_to_docx

    pdf_path = tmp_path / f"{name}.pdf"
    doc.save(str(pdf_path))
    out = tmp_path / f"{name}.docx"
    report = convert_pdf_to_docx(pdf_path.read_bytes(), out, plan=plan)
    return out, report


def _hard_page_breaks(path) -> int:
    """Compte les paragraphes qui commencent RÉELLEMENT sur une page neuve.

    Deux encodages valent saut en Word : ``w:pageBreakBefore`` posé sur le
    paragraphe, ou hérité de son style — c'est ainsi que les styles
    ``Chapitre``/``Partie`` produisent les seuls sauts d'avant ce lot. Ne
    compter qu'un des deux donnerait une mesure fausse, et c'est la mesure qui
    tranche ici.
    """
    import docx
    from docx.oxml.ns import qn

    document = docx.Document(str(path))
    breaking_styles = {
        st.name
        for st in document.styles
        if (ppr := st.element.find(qn("w:pPr"))) is not None
        and ppr.find(qn("w:pageBreakBefore")) is not None
    }

    count = 0
    for para in document.paragraphs:
        if not para.text.strip():
            continue
        ppr = para._p.find(qn("w:pPr"))
        node = None if ppr is None else ppr.find(qn("w:pageBreakBefore"))
        if node is not None:
            # `w:val="0"` DÉSACTIVE le saut du style pour ce paragraphe. Le
            # lire comme un saut inverserait la mesure là où c'est le plus
            # gênant : le tout premier paragraphe du document.
            if node.get(qn("w:val")) not in ("0", "false"):
                count += 1
        elif para.style.name in breaking_styles:
            count += 1
    return count


def _texts(path):
    import docx

    return [p.text for p in docx.Document(str(path)).paragraphs if p.text.strip()]


# ─────────────────────────────────────────────────────────────────────
# Le saut de page — LE point du lot
# ─────────────────────────────────────────────────────────────────────


def test_a_page_break_asked_by_the_model_is_materialized(tmp_path):
    """Franck : « les sauts de page du PDF source sont respectés ». Rien dans
    la chaîne ne pouvait produire cette décision avant ce lot."""
    from app.services.pdf_layout import parse_conversion_plan

    doc = fitz.open()
    _page(doc, [("Dedicace isolee sur sa page.", 0, _BODY_SIZE, False)])
    _page(doc, [("Exergue isole sur sa page.", 0, _BODY_SIZE, False)])

    plan = parse_conversion_plan("p2b0 CORPS SAUT\n")
    out, report = _convert(tmp_path, doc, plan=plan)

    assert _hard_page_breaks(out) == 1
    assert report.page_breaks == 1


def test_the_very_first_paragraph_never_gets_a_page_break(tmp_path):
    """Mesuré sur le manuscrit : un modèle qui statue page par page met SAUT
    sur la page 1 aussi — c'est cohérent de son point de vue. Appliqué tel
    quel, Word ouvre le document sur une page blanche."""
    from app.services.pdf_layout import parse_conversion_plan

    doc = fitz.open()
    _page(doc, [("Premiere page du manuscrit.", 0, _BODY_SIZE, False)])
    _page(doc, [("Deuxieme page du manuscrit.", 0, _BODY_SIZE, False)])

    plan = parse_conversion_plan("p1b0 CORPS SAUT\np2b0 CORPS SAUT\n")
    out, report = _convert(tmp_path, doc, plan=plan)

    assert _hard_page_breaks(out) == 1, "le saut de la page 1 a été appliqué"
    assert report.page_breaks == 1


def test_without_a_plan_the_document_is_unchanged(tmp_path):
    """Échouer OUVERT. Modèle indisponible, plan illisible, quota épuisé : on
    rend ce qu'on rendait avant, jamais un document abîmé."""
    doc = fitz.open()
    _page(doc, [("Dedicace isolee sur sa page.", 0, _BODY_SIZE, False)])
    _page(doc, [("Exergue isole sur sa page.", 0, _BODY_SIZE, False)])

    out, report = _convert(tmp_path, doc, plan=None)

    assert _hard_page_breaks(out) == 0
    assert report.page_breaks == 0
    assert report.missing_chars == 0


def test_an_empty_plan_is_as_good_as_no_plan(tmp_path):
    """Un modèle qui répond à côté ne doit pas valoir mieux ni pire que pas de
    modèle du tout."""
    from app.services.pdf_layout import parse_conversion_plan

    doc = fitz.open()
    _page(doc, [("Dedicace isolee sur sa page.", 0, _BODY_SIZE, False)])
    _page(doc, [("Exergue isole sur sa page.", 0, _BODY_SIZE, False)])

    plan = parse_conversion_plan("Je n'ai pas compris la demande.")
    out, report = _convert(tmp_path, doc, plan=plan)

    assert _hard_page_breaks(out) == 0
    assert report.missing_chars == 0


# ─────────────────────────────────────────────────────────────────────
# L'artefact — le folio resté en texte courant
# ─────────────────────────────────────────────────────────────────────


def test_a_block_marked_artefact_leaves_the_document(tmp_path):
    """Sur le manuscrit de Franck, chaque page finit par son numéro en texte
    courant : 395 « 1 », « 2 », « 3 » au fil du roman."""
    from app.services.pdf_layout import parse_conversion_plan

    doc = fitz.open()
    _page(doc, _body(4) + [("42", _PARA_GAP * 2, _BODY_SIZE, False)])

    plan = parse_conversion_plan("p1b1 ARTEFACT\n")
    out, report = _convert(tmp_path, doc, plan=plan)

    assert "42" not in _texts(out)
    assert report.dropped_blocks == 1


def test_a_dropped_artefact_is_counted_apart_from_a_loss(tmp_path):
    """Sans cette distinction, retirer 395 folios ferait crier le contrôle 0 à
    la perte de données — et le seul signal d'intégrité du convertisseur
    deviendrait un bruit qu'on apprend à ignorer."""
    from app.services.pdf_layout import parse_conversion_plan

    doc = fitz.open()
    _page(doc, _body(4) + [("42", _PARA_GAP * 2, _BODY_SIZE, False)])

    plan = parse_conversion_plan("p1b1 ARTEFACT\n")
    _, report = _convert(tmp_path, doc, plan=plan)

    assert report.missing_chars == 0, "une suppression VOULUE n'est pas une perte"
    assert report.dropped_chars == 2, "les caractères retirés doivent être chiffrés"


# ─────────────────────────────────────────────────────────────────────
# La suite — un paragraphe à cheval sur deux pages
# ─────────────────────────────────────────────────────────────────────


def test_a_continuation_rejoins_the_paragraph_across_the_page(tmp_path):
    """L'inverse du saut, et l'erreur la plus visible quand on se trompe."""
    from app.services.pdf_layout import parse_conversion_plan

    doc = fitz.open()
    _page(doc, [("Le vent se levait sur la ville et personne ne semblait",
                 0, _BODY_SIZE, False)])
    _page(doc, [("y prendre garde ce soir la.", 0, _BODY_SIZE, False)])

    plan = parse_conversion_plan("p2b0 CORPS SUITE\n")
    out, _ = _convert(tmp_path, doc, plan=plan)

    joined = _texts(out)
    assert len(joined) == 1, f"le paragraphe est resté coupé : {joined}"
    assert "semblait y prendre garde" in joined[0]


# ─────────────────────────────────────────────────────────────────────
# Les titres — le modèle tranche, Ely applique
# ─────────────────────────────────────────────────────────────────────


def test_the_model_can_promote_a_block_the_heuristic_missed(tmp_path):
    """Le croisement de signaux d'Ely rate les titres muets et se trompe sur
    les blocs de copyright. C'est exactement le jugement qu'on délègue."""
    from app.services.pdf_layout import parse_conversion_plan

    doc = fitz.open()
    _page(doc, [("Le silence", 0, _BODY_SIZE, False)] + _body(3))

    plan = parse_conversion_plan("p1b0 TITRE1\n")
    out, report = _convert(tmp_path, doc, plan=plan)

    import docx

    styles = [p.style.name for p in docx.Document(str(out)).paragraphs
              if p.text.strip()]
    assert styles[0] == "Chapitre"
    assert report.titles == 1


def test_the_model_can_demote_a_block_the_heuristic_wrongly_promoted(tmp_path):
    """Mesuré sur le manuscrit : le bloc « © S'envoler la mémoire du vent »
    part en style Chapitre, donc avec un saut de page, parce qu'il est haut et
    court. Le modèle doit pouvoir dire non."""
    from app.services.pdf_layout import parse_conversion_plan

    doc = fitz.open()
    _page(doc, [("TITRE EN GRAS BIEN HAUT", 0, 22.0, True)] + _body(3))

    out_before, report_before = _convert(tmp_path, doc, name="avant")
    assert report_before.titles == 1, "le pin suppose que l'heuristique promeut ce bloc"

    plan = parse_conversion_plan("p1b0 CORPS\n")
    out_after, report_after = _convert(tmp_path, doc, plan=plan, name="apres")

    assert report_after.titles == 0
    assert _hard_page_breaks(out_after) == 0


def test_a_subtitle_does_not_open_a_new_page(tmp_path):
    """Mesuré sur le manuscrit : « S'envoler » puis « La mémoire du vent » —
    le modèle rend justement TITRE1 puis TITRE2. Envoyer les deux sur le style
    ``Chapitre`` coupe la page de couverture en deux."""
    from app.services.pdf_layout import parse_conversion_plan

    doc = fitz.open()
    _page(doc, [
        ("Le titre du livre", 0, 22.0, True),
        ("Le sous-titre du livre", _PARA_GAP, 16.0, False),
    ])

    plan = parse_conversion_plan("p1b0 TITRE1\np1b1 TITRE2\n")
    out, report = _convert(tmp_path, doc, plan=plan)

    import docx

    paras = [p for p in docx.Document(str(out)).paragraphs if p.text.strip()]
    assert paras[1].style.name != paras[0].style.name, (
        "TITRE1 et TITRE2 tombent sur le même style"
    )
    assert _hard_page_breaks(out) == 0, "le sous-titre a ouvert une page"
    assert report.titles == 2


def test_a_block_marked_centre_is_actually_centered(tmp_path):
    """Le contrat expose CENTRE au modèle : le lire sans l'appliquer serait
    une promesse tenue à moitié — et invisible, donc pire."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    from app.services.pdf_layout import parse_conversion_plan

    doc = fitz.open()
    _page(doc, [("Une dedicace centree sur sa page.", 0, _BODY_SIZE, False)])

    plan = parse_conversion_plan("p1b0 CORPS CENTRE\n")
    out, _ = _convert(tmp_path, doc, plan=plan)

    import docx

    para = [p for p in docx.Document(str(out)).paragraphs if p.text.strip()][0]
    assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER


# ─────────────────────────────────────────────────────────────────────
# L'intégrité, toujours
# ─────────────────────────────────────────────────────────────────────


def test_the_text_never_comes_from_the_model(tmp_path):
    """Le balisage ne porte que des identifiants. Même si le modèle écrivait du
    texte à côté, aucun caractère ne peut entrer dans le document par là."""
    from app.services.pdf_layout import parse_conversion_plan

    doc = fitz.open()
    _page(doc, _body(4, prefix="Original"))

    plan = parse_conversion_plan(
        "p1b0 CORPS\n"
        "Note du modele : j'ai reformule le paragraphe en « Texte invente ».\n"
    )
    out, report = _convert(tmp_path, doc, plan=plan)

    assert "Texte invente" not in "\n".join(_texts(out))
    assert "Original 0 du paragraphe" in "\n".join(_texts(out))
    assert report.missing_chars == 0
