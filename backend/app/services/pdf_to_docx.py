# =============================================================================
# @project    ELY — Exactly Like You
# @file       backend/app/services/pdf_to_docx.py
# @brief      Reconstruit la structure LOGIQUE d'un PDF, puis l'écrit en
#             styles Word natifs.
#
# @author     Franck OLLIVIER <contact@agent-ely.fr>
# @copyright  Copyright (c) 2025-2026 Franck OLLIVIER
# @license    MIT
#            https://opensource.org/licenses/MIT
# =============================================================================
"""Conversion PDF → DOCX par reconstruction de la structure.

**Ce que faisait Ely avant.** Le cœur tenait en trois lignes :

    for line in text.split("\\n"):
        if line.strip():
            document.add_paragraph(line.strip())

Une ligne VISUELLE = un paragraphe Word. C'est le premier des deux raccourcis
qui condamnent une conversion : on obtient des milliers de faux paragraphes,
des mots coupés par les césures et aucun style. Et ``pypdf.extract_text()`` ne
donne aucune géométrie — ni coordonnées, ni corps, ni graisse — donc rien de
tout cela n'était même détectable.

**Le principe.** Un PDF ne contient pas de paragraphes : il contient des
glyphes posés à des coordonnées. La seule voie utile pour un texte suivi est de
**reconstruire la structure logique depuis la géométrie**, puis de la réécrire
en styles Word natifs.

**Le calibrage n'est pas optionnel.** Les seuils ne sont pas des valeurs
universelles : ce sont les mesures de CE document. Ils sont relevés sur le PDF
lui-même — les avances verticales entre lignes forment deux pics, l'interligne
courant et le passage au paragraphe suivant. On ne les devine pas.

**Périmètre assumé.** Ce module traite : fragments de ligne, césures,
paragraphes à cheval sur deux pages, détection de titre, styles nommés, et un
contrôle d'intégrité des caractères. Il ne traite PAS encore : vers et blocs
centrés, images en ligne, harmonisation de la ponctuation en fausse fonte,
champs PAGEREF du sommaire. Ce sont les raffinements — la structure d'abord.
"""
from __future__ import annotations

import logging
import re
import statistics
from collections import Counter
from dataclasses import dataclass, field, replace
from pathlib import Path

logger = logging.getLogger(__name__)


class NoTextLayer(Exception):
    """Aucune page ne porte de texte — c'est un scan.

    Mieux vaut une erreur explicite qu'un document Word vide livré comme un
    succès : c'est exactement le genre de façade que ce projet traque.
    """


@dataclass
class Calibration:
    """Les mesures de CE document, pas des constantes universelles."""

    line_pitch: float
    para_gap_min: float
    left: float
    body_size: float
    right_edge: float
    body_font: str
    page_width: float
    page_height: float


@dataclass
class ConversionReport:
    """Ce qui a été mesuré — pas seulement « converti »."""

    pages: int = 0
    paragraphs: int = 0
    titles: int = 0
    chars_pdf: int = 0
    chars_docx: int = 0
    missing_chars: int = 0
    calibration: Calibration | None = None
    # Ce que le plan du modèle a produit. Séparé du reste parce que c'est ce
    # que la boucle de conformité doit pouvoir confronter à la demande : « les
    # sauts de page du PDF source sont respectés » se vérifie ici.
    page_breaks: int = 0
    dropped_blocks: int = 0
    # Retirés à la demande du modèle (folios, en-têtes courants). Comptés à
    # part des caractères PERDUS : sans cette séparation, retirer 395 numéros
    # de page ferait crier le contrôle 0 à la perte de données, et le seul
    # signal d'intégrité du convertisseur deviendrait un bruit permanent.
    dropped_chars: int = 0
    plan_applied: bool = False


@dataclass
class _Span:
    """Un fragment homogène : même police, même corps, même style.

    C'est CE niveau que la première version lisait pour détecter les titres,
    puis jetait en écrivant. Les italiques, les changements de corps et les
    couleurs disparaissaient alors que la donnée était sous la main.
    """

    text: str
    size: float
    bold: bool
    italic: bool
    color: int
    font: str


@dataclass
class _Line:
    spans: list[_Span]
    y: float
    x0: float
    x1: float
    page: int

    @property
    def text(self) -> str:
        return "".join(sp.text for sp in self.spans)

    @property
    def size(self) -> float:
        return max((sp.size for sp in self.spans), default=0.0)

    @property
    def bold(self) -> bool:
        return all(sp.bold for sp in self.spans) if self.spans else False


@dataclass
class _Block:
    kind: str                       # "body" | "title"
    lines: list[_Line] = field(default_factory=list)
    # Décidé par le modèle, jamais deviné : ce bloc ouvre-t-il une page ?
    page_break: bool = False
    centered: bool = False
    # Style Word explicite. ``None`` = le défaut du ``kind``. Un titre et un
    # SOUS-titre ne peuvent pas partager un style : le premier ouvre une page,
    # le second non.
    style: str | None = None


# ------------------------------------------------------------------ #
# 1. Calibrage                                                        #
# ------------------------------------------------------------------ #

# Le seuil de paragraphe se place JUSTE AU-DESSUS de l'interligne, pas au
# milieu entre les deux pics : les paragraphes serrés sont plus fréquents que
# les interlignes dilatés, donc mieux vaut pencher du côté de l'interligne.
_PARA_GAP_RATIO = 1.25
# Tolérance verticale pour regrouper les fragments d'une même ligne. Un
# changement d'italique ou une ponctuation en autre fonte crée un bloc
# distinct, parfois décalé de 1 à 2 points.
_LINE_TOLERANCE = 5.0
_DEFAULT_PITCH = 14.0
# En deçà, l'échantillon ne dit rien de fiable sur l'interligne du document.
_MIN_ADVANCE_SAMPLES = 5


def calibrate(doc) -> Calibration:
    """Relève l'interligne, le seuil de paragraphe et la marge sur le document.

    Les avances verticales entre lignes forment deux pics : l'interligne
    courant et le passage au paragraphe suivant. Le mode donne le premier.
    """
    advances: Counter = Counter()
    lefts: Counter = Counter()
    sizes: Counter = Counter()
    fonts: Counter = Counter()
    rights: list[float] = []

    sample = range(min(len(doc), 40))
    for pno in sample:
        page = doc[pno]
        ys = []
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                bbox = line["bbox"]
                ys.append(round(bbox[1], 1))
                lefts[round(bbox[0])] += 1
                rights.append(bbox[2])
                for span in line.get("spans", []):
                    sizes[round(span.get("size", 0), 1)] += 1
                    fonts[str(span.get("font", ""))] += 1
        ys = sorted(set(ys))
        advances.update(round(b - a, 1) for a, b in zip(ys, ys[1:]) if 0 < b - a < 60)

    # Un document de deux lignes donnerait un « pas » égal à l'unique avance
    # mesurée — donc un seuil de paragraphe absurde qui recollerait tout. En
    # dessous de ce nombre d'échantillons, le défaut prudent vaut mieux qu'une
    # mesure qui n'en est pas une.
    pitch = (
        advances.most_common(1)[0][0]
        if sum(advances.values()) >= _MIN_ADVANCE_SAMPLES
        else _DEFAULT_PITCH
    )
    left = float(lefts.most_common(1)[0][0]) if lefts else 72.0
    body = float(sizes.most_common(1)[0][0]) if sizes else 11.0
    # Bord droit du bloc de justification : au-delà, une ligne est « pleine ».
    right = statistics.median(rights) if rights else 500.0

    font = _word_font_name(fonts.most_common(1)[0][0]) if fonts else ""
    rect = doc[0].rect if len(doc) else None

    return Calibration(
        line_pitch=float(pitch),
        para_gap_min=float(pitch) * _PARA_GAP_RATIO,
        left=left,
        body_size=body,
        right_edge=float(right),
        body_font=font,
        page_width=float(rect.width) if rect else 595.0,
        page_height=float(rect.height) if rect else 842.0,
    )


def _word_font_name(pdf_font: str) -> str:
    """Traduit un nom de police PDF en nom de famille lisible par Word.

    Les PDF portent des noms comme ``ABCDEF+Times-Roman`` ou
    ``PublicoText-Italic`` : préfixe de sous-ensemble, et variante accolée. Word
    veut la FAMILLE — il applique ensuite gras et italique par ses propres
    attributs. Sans ce nettoyage, il ne reconnaît rien et substitue sa police
    par défaut, ce qui revient au point de départ.
    """
    name = str(pdf_font or "")
    if "+" in name:                     # préfixe de sous-ensemble
        name = name.split("+", 1)[1]
    for sep in ("-", ","):
        if sep in name:
            name = name.split(sep, 1)[0]
    return {"Times": "Times New Roman", "TimesNewRomanPSMT": "Times New Roman",
            "Helvetica": "Helvetica", "Arial": "Arial"}.get(name, name)


# ------------------------------------------------------------------ #
# 2. Lecture des lignes — piège (a), fragments                        #
# ------------------------------------------------------------------ #

def _is_bold(span: dict) -> bool:
    if span.get("flags", 0) & (1 << 4):
        return True
    return "bold" in str(span.get("font", "")).lower() or "hebo" in str(
        span.get("font", "")
    ).lower()


def _page_lines(page, pno: int) -> list[_Line]:
    """Regroupe les fragments par tolérance verticale, puis trie par ``x``.

    Trié naïvement par ``y``, un texte dont l'italique crée un bloc décalé de
    1 à 2 points part dans le désordre.
    """
    frags: list[tuple[float, float, float, list[_Span]]] = []
    for block in page.get_text("dict")["blocks"]:
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            bbox = line["bbox"]
            spans = [
                _Span(
                    text=sp.get("text", ""),
                    size=round(float(sp.get("size", 0.0)), 1),
                    bold=_is_bold(sp),
                    # Bit 1 des drapeaux PyMuPDF = italique. C'est ce qui
                    # portait 64 spans du manuscrit et disparaissait.
                    italic=bool(sp.get("flags", 0) & (1 << 1)),
                    color=int(sp.get("color", 0) or 0),
                    font=_word_font_name(sp.get("font", "")),
                )
                for sp in line.get("spans", [])
                if sp.get("text")
            ]
            if not any(sp.text.strip() for sp in spans):
                continue
            frags.append(((bbox[1] + bbox[3]) / 2, bbox[0], bbox[2], spans))

    frags.sort(key=lambda f: (f[0], f[1]))

    lines: list[_Line] = []
    for centre, x0, x1, spans in frags:
        if lines and abs(centre - lines[-1].y) <= _LINE_TOLERANCE:
            prev = lines[-1]
            # Même ligne visuelle : on insère dans l'ordre des x.
            if x0 >= prev.x1:
                prev.spans.extend(spans)
            else:
                prev.spans[:0] = spans
            prev.x0 = min(prev.x0, x0)
            prev.x1 = max(prev.x1, x1)
        else:
            lines.append(_Line(spans, centre, x0, x1, pno))
    return lines


# ------------------------------------------------------------------ #
# 3. Césures — piège (b)                                              #
# ------------------------------------------------------------------ #

def join_lines(texts: list[str]) -> str:
    """Recolle des lignes de TEXTE en traitant les césures.

    Conservée pour les appels qui n'ont pas besoin de typographie ; la
    conversion, elle, passe par :func:`join_spans` pour ne pas perdre les
    italiques et les changements de corps.
    """
    return "".join(sp.text for sp in join_spans(
        [[_Span(t, 0.0, False, False, 0, "")] for t in texts]
    ))


def join_spans(per_line: list[list[_Span]]) -> list[_Span]:
    """Recolle les lignes d'un paragraphe **en conservant la typographie**.

    Césures : on supprime le trait d'union si la ligne suivante commence par
    une minuscule ou une apostrophe ; sinon on le garde, et SANS espace —
    « porte-Fenêtre » reste un seul mot.
    """
    out: list[_Span] = []
    for i, spans in enumerate(per_line):
        spans = [sp for sp in spans if sp.text]
        if not spans:
            continue
        nxt = ""
        for follow in per_line[i + 1:]:
            joined = "".join(sp.text for sp in follow).lstrip()
            if joined:
                nxt = joined
                break

        line_text = "".join(sp.text for sp in spans).rstrip()
        last = spans[-1]
        if line_text.endswith("-") and nxt:
            if nxt[0].islower() or nxt[0] in "'’":
                # Trait d'union typographique : on le retire du dernier span.
                trimmed = last.text.rstrip()
                spans[-1] = replace(last, text=trimmed[:-1] if trimmed.endswith("-") else trimmed)
            out.extend(spans)
            continue

        spans[-1] = replace(last, text=last.text.rstrip())
        out.extend(spans)
        if i + 1 < len(per_line):
            out.append(_Span(" ", last.size, last.bold, last.italic, last.color, last.font))

    # Fusion des spans voisins de même typographie : un run par changement
    # réel, pas un run par fragment du PDF.
    merged: list[_Span] = []
    for sp in out:
        if merged and (sp.size, sp.bold, sp.italic, sp.color, sp.font) == (
            merged[-1].size, merged[-1].bold, merged[-1].italic,
            merged[-1].color, merged[-1].font,
        ):
            merged[-1] = replace(merged[-1], text=merged[-1].text + sp.text)
        else:
            merged.append(sp)
    if merged:
        merged[0] = replace(merged[0], text=merged[0].text.lstrip())
        merged[-1] = replace(merged[-1], text=merged[-1].text.rstrip())
    return [sp for sp in merged if sp.text]


# ------------------------------------------------------------------ #
# 4. Titres — piège (d)                                               #
# ------------------------------------------------------------------ #

_TOP_OF_PAGE_RATIO = 0.35


def _toc_titles(doc) -> set[str]:
    """Le signet PDF donne la vérité quand il existe."""
    try:
        return {str(entry[1]).strip().lower() for entry in doc.get_toc() or []}
    except Exception:  # noqa: BLE001 — un PDF sans signet est la norme
        return set()


def _looks_like_title(lines: list[_Line], cal: Calibration, page_h: float,
                      toc: set[str]) -> bool:
    """Jamais par le texte. Chercher « Chapitre » ne marche pas — langue,
    chiffres romains, titres muets. On croise plusieurs signaux."""
    if not lines or len(lines) > 2:
        return False
    text = " ".join(line.text for line in lines).strip()
    if text.lower() in toc:
        return True
    bigger = max(line.size for line in lines) > cal.body_size * 1.15
    bold = all(line.bold for line in lines)
    high = min(line.y for line in lines) < page_h * _TOP_OF_PAGE_RATIO
    return (bigger and high) or (bold and bigger) or (bold and high)


# ------------------------------------------------------------------ #
# 5. Découpage en blocs + fusion inter-pages — piège (c)              #
# ------------------------------------------------------------------ #

_SENTENCE_END = tuple(".!?…»\"'")


def _fills_justification(line: _Line, cal: Calibration) -> bool:
    """La ligne va-t-elle jusqu'au bord du bloc de justification ?"""
    return line.x1 >= cal.right_edge - cal.body_size


def _page_groups(doc, cal: Calibration) -> list[list[list[_Line]]]:
    """Les groupes de lignes de chaque page, **avant** toute fusion.

    Un étage à part parce que c'est l'unité que le modèle doit pouvoir
    désigner : une fois deux pages recollées par l'heuristique, plus personne
    ne peut décider qu'elles auraient dû rester séparées. Une page sans texte
    rend une liste vide — jamais un trou, sinon les numéros de page se
    décalent et les identifiants du balisage désignent le mauvais bloc.
    """
    per_page: list[list[list[_Line]]] = []
    for pno in range(len(doc)):
        lines = _page_lines(doc[pno], pno)
        if not lines:
            per_page.append([])
            continue

        # Découpage vertical : une avance supérieure au seuil ouvre un bloc.
        groups: list[list[_Line]] = [[lines[0]]]
        for prev, cur in zip(lines, lines[1:]):
            if cur.y - prev.y >= cal.para_gap_min:
                groups.append([cur])
            else:
                groups[-1].append(cur)
        per_page.append(groups)
    return per_page


# Ce que chaque rôle du balisage devient dans le document : ``(kind, style)``.
# Une table plutôt qu'une cascade de `if` : ajouter un rôle au contrat doit
# être UNE ligne.
#
# ``TITRE2`` a son propre style, sans saut de page. Mesuré sur le manuscrit :
# le modèle rend justement « S'envoler » en TITRE1 et « La mémoire du vent » en
# TITRE2 — les envoyer tous deux sur `Chapitre` coupait la couverture en deux.
_ROLE_TO_KIND: dict[str, tuple[str, str | None]] = {
    "TITRE1": ("title", "Chapitre"),
    "TITRE2": ("title", "Sous-titre"),
    "CORPS": ("body", None),
    "CITATION": ("body", None),
}


def _build_blocks(doc, cal: Calibration, plan=None) -> tuple[list[_Block], list[list[_Line]]]:
    """Construit les blocs du document, en laissant le plan trancher.

    Un bloc dont le plan ne dit rien garde intégralement le comportement
    d'avant ce lot — heuristique de titre comprise. C'est la règle « échouer
    OUVERT » : le modèle AMÉLIORE un défaut, il ne le remplace pas. Sans plan,
    cette fonction est exactement celle d'avant.

    Returns:
        ``(blocs, groupes retirés)`` — les seconds sont les blocs que le modèle
        a désignés comme artefacts (folios, en-têtes courants). Ils sortent du
        document mais restent comptés, pour que le contrôle d'intégrité
        distingue une suppression VOULUE d'une perte.
    """
    toc = _toc_titles(doc)
    blocks: list[_Block] = []
    dropped: list[list[_Line]] = []

    for pno, groups in enumerate(_page_groups(doc, cal)):
        if not groups:
            continue
        page_h = doc[pno].rect.height

        for gi, group in enumerate(groups):
            block_id = f"p{pno + 1}b{gi}"
            role = plan.role_of(block_id) if plan else None
            # « Le modèle a-t-il statué sur CE bloc ? » — pas « y a-t-il un
            # plan ». Un plan partiel doit laisser les autres blocs tranquilles.
            decided = bool(plan and block_id in plan.decisions)

            if role == "ARTEFACT":
                dropped.append(group)
                continue

            kind, style = _ROLE_TO_KIND.get(role or "") or (
                ("title" if _looks_like_title(group, cal, page_h, toc) else "body"),
                None,
            )

            # Fusion inter-pages. Quand le modèle a statué, sa décision fait
            # foi dans les deux sens : SUITE recolle, son absence sépare.
            merged = False
            can_merge = kind == "body" and blocks and blocks[-1].kind == "body"
            if decided:
                if plan.continues(block_id) and can_merge:
                    blocks[-1].lines.extend(group)
                    merged = True
            elif gi == 0 and can_merge:
                # L'heuristique d'avant : uniquement le PREMIER groupe d'une
                # page, et seulement si la précédente s'arrêtait en pleine phrase.
                last = blocks[-1].lines[-1]
                if (
                    last.page == pno - 1
                    and _fills_justification(last, cal)
                    and not last.text.rstrip().endswith(_SENTENCE_END)
                ):
                    blocks[-1].lines.extend(group)
                    merged = True

            if not merged:
                blocks.append(_Block(
                    kind=kind,
                    lines=list(group),
                    page_break=bool(plan and plan.breaks_before(block_id)),
                    centered=bool(plan and plan.is_centered(block_id)),
                    style=style,
                ))

    return blocks, dropped


# ------------------------------------------------------------------ #
# 6. Écriture DOCX — styles NOMMÉS, pas de formatage direct            #
# ------------------------------------------------------------------ #

def _ensure_styles(document, cal: Calibration) -> None:
    """Pose le format de page et les styles nommés, mesurés sur la source.

    Sans ça, ``python-docx`` produit du US Letter en police par défaut : un
    livre de 396 × 612 ressort en 612 × 792, et le manuscrit ne se ressemble
    plus dès la première page.
    """
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    # Format de page et marges, repris du PDF.
    section = document.sections[0]
    section.page_width = Pt(cal.page_width)
    section.page_height = Pt(cal.page_height)
    margin = Pt(max(cal.left, 28.0))
    section.left_margin = section.right_margin = margin
    section.top_margin = section.bottom_margin = Pt(max(cal.left * 0.9, 28.0))

    # Le corps du roman vit dans `Normal` : le changer est UNE modification.
    normal = document.styles["Normal"]
    if cal.body_font:
        normal.font.name = cal.body_font
    if cal.body_size:
        normal.font.size = Pt(cal.body_size)

    styles = document.styles
    existing = {st.name for st in styles}
    # `Sous-titre` n'ouvre PAS de page : un sous-titre appartient à la page de
    # son titre. C'est la seule différence, et elle est structurante.
    for name, factor, level, breaks in (
        ("Chapitre", 1.5, 0, True),
        ("Partie", 1.9, 0, True),
        ("Sous-titre", 1.2, 1, False),
    ):
        if name in existing:
            continue
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.size = Pt((cal.body_size or 12.0) * factor)
        style.font.bold = True
        ppr = style.element.get_or_add_pPr()
        # `w:outlineLvl` alimente le volet Navigation de Word et les signets à
        # l'export PDF. Sans lui, un titre n'est qu'un texte en gras.
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(level))
        ppr.append(outline)
        if breaks:
            ppr.append(OxmlElement("w:pageBreakBefore"))


def _add_runs(paragraph, spans: list[_Span], cal: Calibration) -> None:
    """Écrit les runs en ne posant QUE ce qui diffère du style.

    Reproduire chaque attribut sur chaque run donnerait un DOCX truffé de
    formatage direct — précisément ce qu'on veut éviter. On laisse le style
    `Normal` porter le cas courant, et on ne surcharge que l'exception.
    """
    from docx.shared import Pt, RGBColor

    for sp in spans:
        run = paragraph.add_run(sp.text)
        if sp.italic:
            run.italic = True
        if sp.bold:
            run.bold = True
        if sp.size and cal.body_size and abs(sp.size - cal.body_size) > 0.4:
            run.font.size = Pt(sp.size)
        if sp.font and sp.font != cal.body_font:
            run.font.name = sp.font
        if sp.color:
            run.font.color.rgb = RGBColor(
                (sp.color >> 16) & 0xFF, (sp.color >> 8) & 0xFF, sp.color & 0xFF
            )


def _write_document(blocks: list[_Block], out_path: Path,
                    cal: Calibration) -> tuple[int, int, int]:
    """Écrit les blocs. Renvoie ``(paragraphes, titres, sauts de page)``.

    Le saut est posé en ``page_break_before`` sur le paragraphe, et non par un
    ``w:br`` dans un paragraphe vide : c'est le même mécanisme que celui des
    styles ``Chapitre``/``Partie``, ça n'ajoute aucun paragraphe fantôme au
    compte, et un auteur qui édite le texte ne se retrouve pas avec un saut
    orphelin au milieu d'un chapitre.

    Les sauts comptés sont les sauts EFFECTIFS — ceux posés ici et ceux
    hérités du style. C'est ce nombre-là qui se confronte aux pages du PDF.
    """
    import docx

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = docx.Document()
    _ensure_styles(document, cal)
    # Ces styles portent déjà `w:pageBreakBefore` : un titre ouvre une page
    # sans qu'on ait rien à poser dessus. `Sous-titre` n'y est PAS.
    breaking_styles = {"Chapitre", "Partie"}

    paragraphs = titles = page_breaks = 0
    for block in blocks:
        spans = join_spans([line.spans for line in block.lines])
        if not any(sp.text.strip() for sp in spans):
            continue
        style = block.style or ("Chapitre" if block.kind == "title" else None)
        para = document.add_paragraph(style=style)
        _add_runs(para, spans, cal)
        if block.centered:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Un saut avant le TOUT PREMIER paragraphe ouvre le document sur une
        # page blanche. Un modèle qui statue page par page met légitimement
        # SAUT sur la page 1 aussi : c'est à la matérialisation de le savoir,
        # pas à lui.
        wanted = block.page_break or style in breaking_styles
        breaks = wanted and paragraphs > 0
        if breaks and style not in breaking_styles:
            para.paragraph_format.page_break_before = True
        elif wanted and not breaks and style in breaking_styles:
            # Le saut vient du STYLE : le retirer demande de le surcharger sur
            # ce paragraphe. Sans ça, un manuscrit qui ouvre sur « Chapitre 1 »
            # commence par une page blanche.
            para.paragraph_format.page_break_before = False
        if breaks:
            page_breaks += 1

        if block.kind == "title":
            titles += 1
        paragraphs += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    return paragraphs, titles, page_breaks


# ------------------------------------------------------------------ #
# 7. Contrôle 0 — le PDF face au DOCX                                 #
# ------------------------------------------------------------------ #

def _char_multiset(text: str) -> Counter:
    """Multi-ensemble des caractères, espaces exclus.

    Insensible à l'ordre de lecture des fragments — qui n'est pas significatif
    dans un PDF — mais impitoyable sur la moindre perte.
    """
    return Counter(c for c in text if not c.isspace())


def _docx_text(path: Path) -> str:
    import docx

    return "\n".join(p.text for p in docx.Document(str(path)).paragraphs)


# ------------------------------------------------------------------ #
# Point d'entrée                                                      #
# ------------------------------------------------------------------ #

def convert_pdf_to_docx(pdf_bytes: bytes, out_path: str | Path,
                        plan=None) -> ConversionReport:
    """Convertit *pdf_bytes* en un ``.docx`` structuré et vérifié.

    Args:
        pdf_bytes: le PDF source.
        out_path: où écrire le ``.docx``.
        plan: un :class:`~app.services.pdf_layout.ConversionPlan` — les
            décisions de structure prises par le modèle. ``None`` (le défaut)
            reproduit exactement le comportement d'avant ce lot : c'est ce qui
            rend la délégation faillible sans risque pour le document.

    Lève :class:`NoTextLayer` si aucune page ne porte de texte.
    """
    import fitz

    out_path = Path(out_path)
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    pdf_text = "".join(doc[i].get_text() for i in range(len(doc)))
    if not pdf_text.strip():
        raise NoTextLayer(f"{len(doc)} page(s) sans couche texte")

    cal = calibrate(doc)
    blocks, dropped = _build_blocks(doc, cal, plan=plan)
    paragraphs, titles, page_breaks = _write_document(blocks, out_path, cal)

    pdf_chars = _char_multiset(pdf_text)
    docx_chars = _char_multiset(_docx_text(out_path))
    dropped_chars = _char_multiset(
        "".join(line.text for group in dropped for line in group)
    )
    # Ce qui manque au DOCX. Deux pertes sont VOULUES et ne sont donc pas des
    # anomalies : le trait d'union d'une césure recollée, et les blocs que le
    # modèle a désignés comme artefacts (folios, en-têtes courants).
    missing = pdf_chars - docx_chars - dropped_chars
    missing.pop("-", None)

    return ConversionReport(
        pages=len(doc),
        paragraphs=paragraphs,
        titles=titles,
        chars_pdf=sum(pdf_chars.values()),
        chars_docx=sum(docx_chars.values()),
        missing_chars=sum(missing.values()),
        calibration=cal,
        page_breaks=page_breaks,
        dropped_blocks=len(dropped),
        dropped_chars=sum(dropped_chars.values()),
        plan_applied=bool(plan and not plan.is_empty),
    )
