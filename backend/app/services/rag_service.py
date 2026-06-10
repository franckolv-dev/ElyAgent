# =============================================================================
# @project    ELY — Exactly Like You
# @file       backend/app/services/rag_service.py
# @brief      RAG (Retrieval-Augmented Generation) document pipeline
#
# @author     Franck OLLIVIER <contact@agent-ely.fr>
# @copyright  Copyright (c) 2025-2026 Franck OLLIVIER — All rights reserved
# @license    Elastic License 2.0
#            https://www.elastic.co/licensing/elastic-license
# @version    1.1.0
# @link       https://github.com/franckolv-dev/PhysicalAgent
#
# RÉSUMÉ DES CONDITIONS :
#   - AUTORISÉ : Usage personnel et professionnel interne (gratuit).
#   - AUTORISÉ : Modification et redistribution avec attribution.
#   - INTERDIT : Revente comme SaaS / service managé à des tiers.
#   - INTERDIT : Suppression des notices de copyright ou de licence.
# =============================================================================
"""RAG (Retrieval-Augmented Generation) document pipeline.

Ingests documents (PDF, TXT, MD, CSV, JSON, DOCX) into a dedicated Qdrant
collection ``knowledge``, chunked with overlap for semantic retrieval.

Reuses the same fastembed embedding model (all-MiniLM-L6-v2, 384 dims) and
Qdrant client as ``memory_manager.py``.
"""
from __future__ import annotations

import asyncio
import logging
import time
import uuid
from functools import lru_cache
from pathlib import Path

from app.services.memory_manager import get_memory_manager

logger = logging.getLogger(__name__)

_COLLECTION_KNOWLEDGE = "knowledge"
_VECTOR_DIM = 384  # all-MiniLM-L6-v2

# Chunking parameters
_CHUNK_SIZE = 1800       # ~512 tokens in chars
_CHUNK_OVERLAP = 175     # ~50 tokens in chars

# Separators for recursive text splitting (tried in order)
_SEPARATORS = ["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " "]


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

def _extract_text_pdf(file_path: Path) -> str:
    """Extract text from a PDF file using pypdf."""
    from pypdf import PdfReader
    reader = PdfReader(str(file_path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _extract_text_txt(file_path: Path) -> str:
    """Extract text from plain text / markdown files."""
    return file_path.read_text(encoding="utf-8", errors="replace")


def _extract_text_csv(file_path: Path) -> str:
    """Extract text from CSV — each row becomes a line."""
    import csv
    rows = []
    with open(file_path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(row))
    return "\n".join(rows)


def _extract_text_json(file_path: Path) -> str:
    """Extract text from JSON — pretty-printed dump."""
    import json
    with open(file_path, encoding="utf-8", errors="replace") as f:
        data = json.load(f)
    return json.dumps(data, indent=2, ensure_ascii=False)


def _extract_text_docx(file_path: Path) -> str:
    """Extract text from DOCX if python-docx is available."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError(
            "Le format DOCX necessite la dependance python-docx. "
            "Installez-la avec: pip install python-docx"
        )
    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also include text from tables — often where the payload sits in reports
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def _extract_text_doc(file_path: Path) -> str:
    """Extract text from legacy MS Word .doc via the `antiword` system binary.

    The OLE2 binary format used by Word 97-2003 has no reliable pure-Python
    parser; antiword is a ~200 KB package shipped in the image instead.
    """
    import subprocess
    try:
        result = subprocess.run(
            ["antiword", "-w", "0", str(file_path)],
            capture_output=True, text=True, timeout=30,
        )
    except FileNotFoundError:
        raise ValueError(
            "Le format .doc necessite le binaire `antiword`. "
            "Il doit etre installe dans l'image Docker (apt-get install antiword)."
        )
    if result.returncode != 0:
        raise ValueError(
            f"Echec de l'extraction .doc: {result.stderr.strip() or 'erreur inconnue'}. "
            "Si c'est un fichier recent, ouvre-le dans Word et resauvegarde-le en .docx."
        )
    return result.stdout


def _extract_text_xlsx(file_path: Path) -> str:
    """Extract text from .xlsx — one line per row, sheets separated by title."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ValueError(
            "Le format XLSX necessite la dependance openpyxl. "
            "Installez-la avec: pip install openpyxl"
        )
    wb = load_workbook(filename=str(file_path), read_only=True, data_only=True)
    sheets_text = []
    for sheet in wb.worksheets:
        sheet_lines = [f"## {sheet.title}"]
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                sheet_lines.append(" | ".join(cells))
        if len(sheet_lines) > 1:
            sheets_text.append("\n".join(sheet_lines))
    wb.close()
    return "\n\n".join(sheets_text)


def _extract_text_xls(file_path: Path) -> str:
    """Extract text from legacy Excel 97-2003 .xls via xlrd."""
    try:
        import xlrd
    except ImportError:
        raise ValueError(
            "Le format XLS necessite la dependance xlrd. "
            "Installez-la avec: pip install 'xlrd>=2.0.1'"
        )
    wb = xlrd.open_workbook(str(file_path))
    sheets_text = []
    for sheet in wb.sheets():
        sheet_lines = [f"## {sheet.name}"]
        for row_idx in range(sheet.nrows):
            row = sheet.row_values(row_idx)
            cells = [str(c) if c != "" else "" for c in row]
            if any(c.strip() for c in cells):
                sheet_lines.append(" | ".join(cells))
        if len(sheet_lines) > 1:
            sheets_text.append("\n".join(sheet_lines))
    return "\n\n".join(sheets_text)


_EXTRACTORS: dict[str, callable] = {
    ".pdf": _extract_text_pdf,
    ".txt": _extract_text_txt,
    ".md": _extract_text_txt,
    ".csv": _extract_text_csv,
    ".json": _extract_text_json,
    ".doc": _extract_text_doc,
    ".docx": _extract_text_docx,
    ".xls": _extract_text_xls,
    ".xlsx": _extract_text_xlsx,
}


# ---------------------------------------------------------------------------
# Recursive text chunking
# ---------------------------------------------------------------------------

def _split_text(text: str, chunk_size: int = _CHUNK_SIZE, chunk_overlap: int = _CHUNK_OVERLAP) -> list[str]:
    """Split *text* recursively using paragraph/sentence/word boundaries.

    Strategy: try each separator in order. When a chunk exceeds *chunk_size*,
    split at the next separator level. Overlap is added by prepending the tail
    of the previous chunk.
    """
    if len(text) <= chunk_size:
        return [text.strip()] if text.strip() else []

    # Find the best separator that actually occurs in the text
    sep = None
    for s in _SEPARATORS:
        if s in text:
            sep = s
            break
    if sep is None:
        # No separator found — hard split at chunk_size
        return _hard_split(text, chunk_size, chunk_overlap)

    parts = text.split(sep)
    chunks: list[str] = []
    current = ""

    for part in parts:
        candidate = (current + sep + part) if current else part
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            if current.strip():
                chunks.append(current.strip())
            # If the single part is too large, recurse with next separator
            if len(part) > chunk_size:
                sub_chunks = _split_text(part, chunk_size, chunk_overlap)
                chunks.extend(sub_chunks)
                current = ""
            else:
                current = part

    if current.strip():
        chunks.append(current.strip())

    # Add overlap between consecutive chunks
    if chunk_overlap > 0 and len(chunks) > 1:
        chunks = _add_overlap(chunks, chunk_overlap)

    return chunks


def _hard_split(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    """Fallback: split text at fixed character boundaries."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - chunk_overlap if chunk_overlap < chunk_size else end
    return chunks


def _add_overlap(chunks: list[str], overlap: int) -> list[str]:
    """Prepend the tail of the previous chunk to each subsequent chunk."""
    result = [chunks[0]]
    for i in range(1, len(chunks)):
        prev_tail = chunks[i - 1][-overlap:]
        # Only add overlap if it doesn't start mid-word
        space_idx = prev_tail.find(" ")
        if space_idx > 0:
            prev_tail = prev_tail[space_idx + 1:]
        combined = prev_tail + " " + chunks[i]
        result.append(combined.strip())
    return result


# ---------------------------------------------------------------------------
# RAG Service
# ---------------------------------------------------------------------------

class RAGService:
    """Document ingestion, retrieval, and management backed by Qdrant."""

    def __init__(self) -> None:
        self._mm = get_memory_manager()

    @property
    def client(self):
        return self._mm.client

    @property
    def encoder(self):
        return self._mm.encoder

    async def _embed(self, text: str) -> list[float]:
        """Reuse the memory_manager embedding with its LRU cache."""
        return await self._mm._embed(text)

    # ------------------------------------------------------------------
    # Collection initialisation
    # ------------------------------------------------------------------

    async def init_collection(self) -> None:
        """Create the 'knowledge' collection in Qdrant if it does not exist."""
        try:
            from qdrant_client.models import Distance, VectorParams
            collections_resp = await asyncio.to_thread(self.client.get_collections)
            existing = {c.name for c in collections_resp.collections}
            if _COLLECTION_KNOWLEDGE not in existing:
                await asyncio.to_thread(
                    self.client.create_collection,
                    _COLLECTION_KNOWLEDGE,
                    vectors_config=VectorParams(
                        size=_VECTOR_DIM, distance=Distance.COSINE
                    ),
                )
                logger.info("Qdrant collection '%s' created", _COLLECTION_KNOWLEDGE)
            else:
                logger.info("Qdrant collection '%s' already exists", _COLLECTION_KNOWLEDGE)
            # Index payload user_id (revue 2026-06-10 §4) — idempotent,
            # appliqué aussi à une collection existante.
            try:
                from qdrant_client.models import PayloadSchemaType
                await asyncio.to_thread(
                    self.client.create_payload_index,
                    _COLLECTION_KNOWLEDGE,
                    field_name="user_id",
                    field_schema=PayloadSchemaType.KEYWORD,
                )
            except Exception:
                pass
        except Exception as exc:
            logger.warning("Failed to init knowledge collection: %s", exc)

    # ------------------------------------------------------------------
    # Ingestion
    # ------------------------------------------------------------------

    async def ingest_document(
        self,
        file_path: str | Path,
        user_id: str,
        title: str | None = None,
        collection_name: str | None = None,
        source_file: str | None = None,
    ) -> dict:
        """Chunk, embed, and store a document in Qdrant.

        Args:
            source_file: optional override for the ``source_file`` payload
                stored alongside each chunk. Used by the auto-indexer to
                preserve the user-facing original path (e.g.
                ``/Users/franck/Documents/budget.xlsx``) instead of the
                temp file path used during ingestion.

        Returns:
            dict with document_id, title, chunk_count, status.
        """
        file_path = Path(file_path)
        collection = collection_name or _COLLECTION_KNOWLEDGE
        ext = file_path.suffix.lower()

        if ext not in _EXTRACTORS:
            raise ValueError(f"Type de fichier non supporte pour l'indexation : {ext}")

        # Extract text
        try:
            text = await asyncio.to_thread(_EXTRACTORS[ext], file_path)
        except Exception as exc:
            raise ValueError(f"Erreur lors de l'extraction du texte : {exc}") from exc

        if not text or not text.strip():
            raise ValueError("Le document est vide ou ne contient pas de texte extractible.")

        # Generate document ID and title
        document_id = str(uuid.uuid4())
        doc_title = title or file_path.stem
        # Honour the explicit override if the caller (auto-indexer) passed it,
        # so the chunk payload stores the original user-facing path instead
        # of the temp file used for extraction.
        source_file = source_file or file_path.name

        # Chunk the text
        chunks = _split_text(text)
        if not chunks:
            raise ValueError("Aucun chunk de texte genere apres le decoupage du document.")

        total_chunks = len(chunks)
        logger.info(
            "Ingesting '%s' (%d chunks) for user %s as doc %s",
            source_file, total_chunks, user_id, document_id,
        )

        # Embed and store each chunk
        from qdrant_client.models import PointStruct

        points = []
        for idx, chunk_text in enumerate(chunks):
            vector = await self._embed(chunk_text)
            point_id = str(uuid.uuid4())
            payload = {
                "content": chunk_text,
                "user_id": user_id,
                "document_id": document_id,
                "title": doc_title,
                "source_file": source_file,
                "chunk_index": idx,
                "total_chunks": total_chunks,
                "created_at": time.time(),
            }
            points.append(PointStruct(id=point_id, vector=vector, payload=payload))

        # Batch upsert
        await asyncio.to_thread(
            self.client.upsert,
            collection_name=collection,
            points=points,
        )

        logger.info("Document '%s' ingested: %d chunks stored", source_file, total_chunks)
        return {
            "document_id": document_id,
            "title": doc_title,
            "source_file": source_file,
            "chunk_count": total_chunks,
            "status": "ingested",
        }

    # ------------------------------------------------------------------
    # Retrieval
    # ------------------------------------------------------------------

    async def search_knowledge(
        self,
        query: str,
        user_id: str,
        limit: int = 5,
        score_threshold: float = 0.4,
    ) -> list[dict]:
        """Search the knowledge base for chunks relevant to *query*.

        Multi-tenant safety: refuse empty user_id — would leak across users
        whose payload also has user_id="". See memory_manager._qdrant_candidates
        for the same guard pattern.

        Returns:
            list of {content, title, source_file, chunk_index, score}
        """
        if not user_id:
            logger.warning("search_knowledge refused: empty user_id (tenant isolation)")
            return []
        try:
            from qdrant_client.models import FieldCondition, Filter, MatchValue

            vector = await self._embed(query)
            result = await asyncio.to_thread(
                self.client.query_points,
                collection_name=_COLLECTION_KNOWLEDGE,
                query=vector,
                query_filter=Filter(
                    must=[FieldCondition(key="user_id", match=MatchValue(value=user_id))]
                ),
                limit=limit,
                score_threshold=score_threshold,
                with_payload=True,
            )
            return [
                {
                    "content": hit.payload.get("content", ""),
                    "title": hit.payload.get("title", ""),
                    "source_file": hit.payload.get("source_file", ""),
                    "chunk_index": hit.payload.get("chunk_index", 0),
                    "total_chunks": hit.payload.get("total_chunks", 0),
                    "score": hit.score,
                }
                for hit in result.points
            ]
        except Exception as exc:
            logger.warning("Knowledge search failed: %s", exc)
            return []

    # ------------------------------------------------------------------
    # Management
    # ------------------------------------------------------------------

    async def list_documents(self, user_id: str) -> list[dict]:
        """List all documents ingested by *user_id*.

        Multi-tenant safety: refuse empty user_id (see search_knowledge).

        Returns a deduplicated list by document_id with metadata.
        """
        if not user_id:
            logger.warning("list_documents refused: empty user_id (tenant isolation)")
            return []
        try:
            from qdrant_client.models import FieldCondition, Filter, MatchValue

            # Scroll all points for this user in the knowledge collection
            all_points = []
            offset = None
            while True:
                scroll_kwargs = {
                    "collection_name": _COLLECTION_KNOWLEDGE,
                    "scroll_filter": Filter(
                        must=[FieldCondition(key="user_id", match=MatchValue(value=user_id))]
                    ),
                    "limit": 100,
                    "with_payload": True,
                }
                if offset is not None:
                    scroll_kwargs["offset"] = offset
                result = await asyncio.to_thread(self.client.scroll, **scroll_kwargs)
                points, next_offset = result
                all_points.extend(points)
                if next_offset is None:
                    break
                offset = next_offset

            # Deduplicate by document_id
            docs: dict[str, dict] = {}
            for p in all_points:
                doc_id = p.payload.get("document_id", "")
                if doc_id not in docs:
                    docs[doc_id] = {
                        "document_id": doc_id,
                        "title": p.payload.get("title", ""),
                        "source_file": p.payload.get("source_file", ""),
                        "chunk_count": p.payload.get("total_chunks", 0),
                        "created_at": p.payload.get("created_at"),
                    }
            return list(docs.values())
        except Exception as exc:
            logger.warning("list_documents failed: %s", exc)
            return []

    async def get_document_info(self, document_id: str, user_id: str) -> dict | None:
        if not user_id:
            logger.warning("get_document_info refused: empty user_id (tenant isolation)")
            return None
        # ── original body below ───────────────────────────────────────────
        """Get metadata for a specific document."""
        try:
            from qdrant_client.models import FieldCondition, Filter, MatchValue

            result = await asyncio.to_thread(
                self.client.scroll,
                collection_name=_COLLECTION_KNOWLEDGE,
                scroll_filter=Filter(
                    must=[
                        FieldCondition(key="user_id", match=MatchValue(value=user_id)),
                        FieldCondition(key="document_id", match=MatchValue(value=document_id)),
                    ]
                ),
                limit=1,
                with_payload=True,
            )
            points, _ = result
            if not points:
                return None
            p = points[0]
            return {
                "document_id": document_id,
                "title": p.payload.get("title", ""),
                "source_file": p.payload.get("source_file", ""),
                "chunk_count": p.payload.get("total_chunks", 0),
                "created_at": p.payload.get("created_at"),
            }
        except Exception as exc:
            logger.warning("get_document_info failed: %s", exc)
            return None

    async def delete_document(self, document_id: str, user_id: str) -> bool:
        if not user_id:
            logger.warning("delete_document refused: empty user_id (tenant isolation)")
            return False
        # ── original body below ───────────────────────────────────────────
        """Delete all chunks for a document from Qdrant.

        Returns True on success, False on failure.
        """
        try:
            from qdrant_client.models import (
                FieldCondition,
                Filter,
                FilterSelector,
                MatchValue,
            )

            # Qdrant supports deleting by filter via FilterSelector
            await asyncio.to_thread(
                self.client.delete,
                collection_name=_COLLECTION_KNOWLEDGE,
                points_selector=FilterSelector(
                    filter=Filter(
                        must=[
                            FieldCondition(key="user_id", match=MatchValue(value=user_id)),
                            FieldCondition(key="document_id", match=MatchValue(value=document_id)),
                        ]
                    )
                ),
            )
            logger.info("Deleted document %s for user %s", document_id, user_id)
            return True
        except Exception as exc:
            logger.warning("delete_document failed: %s", exc)
            return False


@lru_cache(maxsize=1)
def get_rag_service() -> RAGService:
    return RAGService()
